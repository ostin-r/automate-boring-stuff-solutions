'''
Austin Richards 4/5/21

custom_invitation.py creates personalied invitations and puts
them all on separate pages in a .docx file.  Invitations are
personalized by reading a .txt file of recipient names.
'''
import os
import docx
from docx.shared import Pt


def custom_invitation(text_file):
    names_file = open(text_file, 'r')
    names = names_file.readlines()
    names_file.close()
    doc = docx.Document()
    page_break = 4

    for name in names:
        greeting = 'It would be a pleasure to have the company of\n'
        address = 'at 11010 Memory Lane on the evening of\n'
        date = 'April 1st\n'
        time = "at 7 o'clock"

        add_greeting = doc.add_paragraph(greeting, 'Heading 1')
        add_name = doc.add_paragraph(name, 'Intense Quote')
        add_address = doc.add_paragraph(address, 'Heading 1')
        add_date = doc.add_paragraph(date, 'Heading 1')
        add_time = doc.add_paragraph(time, 'Heading 1')

        add_greeting.alignment = 1
        add_name.alignment = 1
        add_address.alignment = 1
        add_date.alignment = 1
        add_time.alignment = 1

        doc.paragraphs[page_break].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)

        page_break += 5
    
    doc.save('invitations.docx')


os.chdir('Chapter 15')
custom_invitation('name_list.txt')